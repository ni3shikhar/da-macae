"""Document Generator — Converts agent output to Word/Excel and uploads to blob storage.

After an agent finishes, this module analyses the final text output and
generates a downloadable Word (.docx) or Excel (.xlsx) document, then
uploads it to Azure Blob Storage (Azurite in dev).  The upload result
JSON is returned so callers can append it to the agent output and the
frontend will render a download card automatically.

Also provides summary generation to show condensed output in chat while
keeping full details in the document.
"""

from __future__ import annotations

import io
import json
import os
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Any

import structlog

logger = structlog.get_logger(__name__)


@dataclass
class DocumentResult:
    """Result of document generation with summary for chat display."""
    summary: str  # Condensed version for chat window
    full_text: str  # Original detailed text (stored in document)
    uploads: list[str]  # Blob upload JSON strings

# ── Minimum content length to trigger document generation ───────────
_MIN_REPORT_LENGTH = 500

# Keywords in agent names or content that signal "report-like" output
_REPORT_AGENT_KEYWORDS = {
    "report", "discovery", "analysis", "quality", "mapping",
    "assessment", "audit", "summary", "migration",
}

# Subtask labels that indicate intermediate/data-gathering steps
# where document generation should be SKIPPED.
_INTERMEDIATE_SUBTASK_KEYWORDS = {
    "list", "retrieve", "fetch", "discover", "connect",
    "initialize", "setup", "configure", "verify", "check",
    "load", "get", "query", "scan", "collect", "gather",
    "identify", "enumerate", "catalog", "catalogue",
}

# Subtask labels that indicate final/report steps where document
# generation SHOULD be triggered.
_REPORT_SUBTASK_KEYWORDS = {
    "report", "compile", "generate report", "produce report",
    "summarize", "summarise", "final", "assessment report",
    "excel report", "comprehensive", "consolidate",
}


def _looks_like_report(
    text: str,
    agent_name: str,
    subtask_label: str | None = None,
) -> bool:
    """Heuristic: does this output look like a report worth exporting?

    Skips document generation for intermediate/data-gathering subtasks
    (e.g. 'List available subscriptions', 'Retrieve MCSB controls').
    Only generates documents for final report or substantial analysis output.
    """
    if len(text) < _MIN_REPORT_LENGTH:
        return False

    # ── Subtask-level gating ───────────────────────────────────────
    if subtask_label:
        label_lower = subtask_label.lower()
        # Explicitly a report subtask → allow
        if any(kw in label_lower for kw in _REPORT_SUBTASK_KEYWORDS):
            return True
        # Explicitly an intermediate subtask → skip
        if any(
            label_lower.startswith(kw) or f" {kw} " in f" {label_lower} "
            for kw in _INTERMEDIATE_SUBTASK_KEYWORDS
        ):
            return False

    # ── Content-based heuristic ────────────────────────────────────
    # Check content for report-like structure (headings, tables, lists)
    heading_count = len(re.findall(r"^#{1,4}\s", text, re.MULTILINE))
    table_count = len(re.findall(r"^\|", text, re.MULTILINE))
    list_count = len(re.findall(r"^[-*]\s", text, re.MULTILINE))
    # Needs substantial structure to qualify as a report
    if heading_count >= 3 and (table_count >= 4 or list_count >= 8):
        return True
    if table_count >= 8:
        return True
    if heading_count >= 4 and list_count >= 10:
        return True
    return False


def _has_significant_tables(text: str) -> bool:
    """Check if the text contains markdown tables that would benefit from Excel."""
    tables = re.findall(
        r"(\|[^\n]+\|\n\|[-:\s|]+\|\n(?:\|[^\n]+\|\n?)+)", text
    )
    if not tables:
        return False
    # At least one table with 3+ data rows
    for tbl in tables:
        data_rows = [
            line for line in tbl.strip().split("\n")
            if line.strip().startswith("|") and not re.match(r"^\|[\s:-]+\|$", line.strip())
        ]
        # Subtract header row
        if len(data_rows) >= 4:  # header + 3 data rows
            return True
    return False


# ── Summary Generation ──────────────────────────────────────────────

def _generate_summary(text: str, agent_name: str, subtask_label: str | None = None) -> str:
    """Generate a condensed summary from detailed agent output for chat display.
    
    Extracts:
    - Main headings (## level)
    - Key findings/metrics (numbers, percentages)
    - Table row counts instead of full tables
    - First sentence of each major section
    """
    lines = text.split("\n")
    summary_parts: list[str] = []
    
    # Title
    title = subtask_label or agent_name
    summary_parts.append(f"## {title} - Summary\n")
    
    # Extract main headings and first meaningful line after each
    current_section: str | None = None
    section_first_line: str | None = None
    tables_found = 0
    table_rows_total = 0
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip blob upload JSON
        if line.startswith("{") and '"status":"uploaded"' in line:
            i += 1
            continue
        
        # Level 2 heading - main section
        if line.startswith("## ") and not line.startswith("###"):
            # Save previous section summary
            if current_section and section_first_line:
                summary_parts.append(f"**{current_section}**: {section_first_line}")
            current_section = line[3:].strip()
            section_first_line = None
            i += 1
            continue
        
        # Capture first meaningful line of section
        if current_section and not section_first_line and line and not line.startswith("#"):
            # Skip table lines, capture text
            if not line.startswith("|"):
                # Truncate long lines
                if len(line) > 150:
                    line = line[:147] + "..."
                section_first_line = line
        
        # Count tables instead of including them
        if line.startswith("|") and i + 1 < len(lines):
            table_start = i
            row_count = 0
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not re.match(r"^\|[\s:-]+\|$", lines[i].strip()):
                    row_count += 1
                i += 1
            tables_found += 1
            table_rows_total += row_count - 1  # Exclude header
            continue
        
        i += 1
    
    # Add last section
    if current_section and section_first_line:
        summary_parts.append(f"**{current_section}**: {section_first_line}")
    
    # Extract key metrics (numbers with context)
    metrics = re.findall(
        r"(?:total|count|found|identified|analyzed|processed|risk|score)[:\s]+(\d+(?:,\d+)*(?:\.\d+)?%?)",
        text,
        re.IGNORECASE,
    )
    if metrics:
        unique_metrics = list(dict.fromkeys(metrics[:5]))  # First 5 unique
        summary_parts.append(f"\n**Key Metrics**: {', '.join(unique_metrics)}")
    
    # Table summary
    if tables_found > 0:
        summary_parts.append(f"\n**Data**: {tables_found} table(s) with {table_rows_total} total rows")
    
    # Footer pointing to document
    summary_parts.append("\n---\n*📄 Full details available in the attached document(s).*")
    
    return "\n".join(summary_parts)


# ── Word Document Generation ────────────────────────────────────────


def _generate_word_doc(text: str, title: str) -> bytes:
    """Convert markdown-ish text into a Word .docx document."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Title ───────────────────────────────────────────────────
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle with timestamp
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        f"Generated on {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # spacer

    # ── Parse markdown line-by-line ─────────────────────────────
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip blob upload JSON lines
        if stripped.startswith("{") and '"status":"uploaded"' in stripped:
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)  # docx supports 0-9
            doc.add_heading(heading_match.group(2), level=level)
            i += 1
            continue

        # Markdown table
        if stripped.startswith("|") and i + 1 < len(lines):
            # Collect all table lines
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            _add_table_to_doc(doc, table_lines)
            continue

        # Bullet / numbered lists
        list_match = re.match(r"^([-*]|\d+\.)\s+(.*)", stripped)
        if list_match:
            para = doc.add_paragraph(list_match.group(2), style="List Bullet")
            i += 1
            continue

        # Bold text lines (e.g., **Section:**)
        if stripped.startswith("**") and stripped.endswith("**"):
            para = doc.add_paragraph()
            run = para.add_run(stripped.strip("*"))
            run.bold = True
            i += 1
            continue

        # JSON block — try to convert to a table
        if stripped.startswith("[") or (stripped.startswith("{") and '"status"' not in stripped[:50]):
            json_lines = [lines[i]]
            bracket_depth = stripped.count("[") + stripped.count("{") - stripped.count("]") - stripped.count("}")
            i += 1
            while i < len(lines) and bracket_depth > 0:
                json_lines.append(lines[i])
                bracket_depth += lines[i].count("[") + lines[i].count("{") - lines[i].count("]") - lines[i].count("}")
                i += 1
            json_text = "\n".join(json_lines)
            if _try_add_json_table(doc, json_text):
                continue
            # Fallback — add as formatted text
            para = doc.add_paragraph(json_text[:2000])
            para.style = doc.styles.get("No Spacing", doc.styles["Normal"])
            continue

        # Regular paragraph (skip empty lines)
        if stripped:
            para = doc.add_paragraph(stripped)
        i += 1

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _try_add_json_table(doc: Any, json_text: str) -> bool:
    """Try to parse JSON and add it as a Word table.

    Handles both JSON arrays of objects and single objects.
    Returns True if successful, False to fall back to text rendering.
    """
    from docx.shared import Pt

    try:
        data = json.loads(json_text)
    except (json.JSONDecodeError, ValueError):
        return False

    # Normalise to list of flat dicts
    rows: list[dict[str, str]] = []
    if isinstance(data, dict):
        # Single object — check for nested list (e.g. {"findings": [...]})
        for v in data.values():
            if isinstance(v, list) and v and isinstance(v[0], dict):
                rows = v
                break
        if not rows:
            rows = [data]
    elif isinstance(data, list) and data and isinstance(data[0], dict):
        rows = data
    else:
        return False

    if not rows:
        return False

    # Pick display columns (skip very long or nested values)
    all_keys = list(dict.fromkeys(k for row in rows[:50] for k in row.keys()))
    display_keys = [
        k for k in all_keys
        if not any(
            isinstance(row.get(k), (dict, list)) and len(str(row.get(k, ""))) > 200
            for row in rows[:5]
        )
    ][:10]  # max 10 columns

    if not display_keys:
        return False

    num_cols = len(display_keys)
    table = doc.add_table(rows=1 + min(len(rows), 100), cols=num_cols)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        pass

    # Header
    for ci, key in enumerate(display_keys):
        cell = table.rows[0].cells[ci]
        cell.text = str(key).replace("_", " ").title()
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)

    # Data rows (cap at 100)
    for ri, row in enumerate(rows[:100]):
        for ci, key in enumerate(display_keys):
            val = row.get(key, "")
            if isinstance(val, (dict, list)):
                val = json.dumps(val, default=str)[:150]
            cell_text = str(val)[:200]
            table.rows[ri + 1].cells[ci].text = cell_text
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(7)

    if len(rows) > 100:
        doc.add_paragraph(f"... and {len(rows) - 100} more rows")

    return True


def _add_table_to_doc(doc: Any, table_lines: list[str]) -> None:
    """Parse a markdown table and add it as a Word table."""
    from docx.shared import Pt, Inches

    if len(table_lines) < 2:
        return

    # Parse header
    header_cells = [c.strip() for c in table_lines[0].split("|") if c.strip()]

    # Skip separator line
    data_start = 1
    if re.match(r"^\|[\s:-]+\|$", table_lines[1]):
        data_start = 2

    # Parse data rows
    data_rows: list[list[str]] = []
    for line in table_lines[data_start:]:
        cells = [c.strip() for c in line.split("|") if c.strip()]
        data_rows.append(cells)

    if not header_cells:
        return

    # Create table
    num_cols = len(header_cells)
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.style = "Light Grid Accent 1"

    # Header row
    for j, cell_text in enumerate(header_cells):
        if j < num_cols:
            cell = table.rows[0].cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

    # Data rows
    for row_idx, row_data in enumerate(data_rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols and (row_idx + 1) < len(table.rows):
                cell = table.rows[row_idx + 1].cells[j]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    doc.add_paragraph()  # spacer after table


# ── Excel Generation ────────────────────────────────────────────────


def _generate_excel(text: str, title: str) -> bytes:
    """Extract markdown tables from text and create an Excel workbook."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    # Remove default sheet
    wb.remove(wb.active)

    # Find all markdown tables with optional preceding heading
    sections = re.split(r"(?=^#{1,4}\s)", text, flags=re.MULTILINE)

    sheet_num = 0
    for section in sections:
        # Find heading for this section
        heading_match = re.match(r"^#{1,4}\s+(.*)", section.strip())
        section_title = heading_match.group(1).strip() if heading_match else ""

        # Find tables in this section
        table_pattern = re.compile(
            r"(\|[^\n]+\|\n\|[-:\s|]+\|\n(?:\|[^\n]+\|\n?)+)"
        )
        for tbl_match in table_pattern.finditer(section):
            table_text = tbl_match.group(1)
            table_lines = [l.strip() for l in table_text.strip().split("\n") if l.strip()]
            if len(table_lines) < 3:
                continue

            sheet_num += 1
            # Sheet name from heading or generic
            sheet_name = section_title[:31] if section_title else f"Table {sheet_num}"
            # Excel sheet names must be unique and <= 31 chars
            sheet_name = re.sub(r"[\\/*?\[\]:]", "", sheet_name)[:31]

            # Ensure unique name
            existing_names = [ws.title for ws in wb.worksheets]
            if sheet_name in existing_names:
                sheet_name = f"{sheet_name[:27]}_{sheet_num}"

            ws = wb.create_worksheet(title=sheet_name)

            # Parse header
            header_cells = [c.strip() for c in table_lines[0].split("|") if c.strip()]

            # Skip separator
            data_start = 1
            if re.match(r"^\|[\s:-]+\|$", table_lines[1]):
                data_start = 2

            # Header styling
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF", size=10)
            thin_border = Border(
                left=Side(style="thin"),
                right=Side(style="thin"),
                top=Side(style="thin"),
                bottom=Side(style="thin"),
            )

            for col_idx, header in enumerate(header_cells, 1):
                cell = ws.cell(row=1, column=col_idx, value=header)
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center")
                cell.border = thin_border

            # Data rows
            for row_idx, line in enumerate(table_lines[data_start:], 2):
                cells = [c.strip() for c in line.split("|") if c.strip()]
                for col_idx, val in enumerate(cells, 1):
                    if col_idx <= len(header_cells):
                        # Try to convert numeric values
                        try:
                            if "." in val:
                                numeric_val = float(val.replace(",", ""))
                                cell = ws.cell(row=row_idx, column=col_idx, value=numeric_val)
                            elif val.isdigit() or (val.startswith("-") and val[1:].isdigit()):
                                cell = ws.cell(row=row_idx, column=col_idx, value=int(val.replace(",", "")))
                            else:
                                cell = ws.cell(row=row_idx, column=col_idx, value=val)
                        except (ValueError, AttributeError):
                            cell = ws.cell(row=row_idx, column=col_idx, value=val)
                        cell.border = thin_border
                        cell.font = Font(size=10)

            # Auto-fit column widths (approximate)
            for col_idx in range(1, len(header_cells) + 1):
                max_len = len(str(header_cells[col_idx - 1]))
                for row in ws.iter_rows(min_row=2, min_col=col_idx, max_col=col_idx):
                    for cell in row:
                        if cell.value:
                            max_len = max(max_len, len(str(cell.value)))
                ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = min(max_len + 4, 50)

            # Freeze header row
            ws.freeze_panes = "A2"

    # If no tables found, create a summary sheet with the text
    if sheet_num == 0:
        ws = wb.create_worksheet(title="Report")
        ws.cell(row=1, column=1, value=title).font = Font(bold=True, size=14)
        for row_idx, line in enumerate(text.split("\n"), 3):
            ws.cell(row=row_idx, column=1, value=line.strip())

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── Blob Upload (direct to Azure Storage) ───────────────────────────


async def _upload_binary_blob(
    container: str, blob_name: str, data: bytes, content_type: str
) -> str:
    """Upload binary data directly to Azure Storage / Azurite.

    Returns the same JSON format as the MCP storage_upload_blob tool
    so the frontend renders a download card.
    """
    from azure.storage.blob.aio import BlobServiceClient
    from azure.storage.blob import ContentSettings

    conn_str = os.getenv("AZURE_STORAGE_CONNECTION_STRING", "")
    if not conn_str:
        raise RuntimeError("AZURE_STORAGE_CONNECTION_STRING not configured")

    async with BlobServiceClient.from_connection_string(conn_str) as bsc:
        cc = bsc.get_container_client(container)
        try:
            await cc.create_container()
        except Exception:
            pass  # already exists
        bc = cc.get_blob_client(blob_name)
        await bc.upload_blob(
            data,
            overwrite=True,
            content_settings=ContentSettings(content_type=content_type),
        )

    return json.dumps({"status": "uploaded", "container": container, "blob": blob_name})


# ── Public API ──────────────────────────────────────────────────────


async def generate_and_upload_documents(
    text: str,
    agent_name: str,
    container: str = "migration-reports",
    subtask_label: str | None = None,
) -> list[str]:
    """Generate Word and/or Excel documents from agent output and upload them.

    Returns a list of blob upload JSON strings (same format as
    ``storage_upload_blob``) that callers can append to agent output so
    the frontend renders download cards.

    - **Word (.docx)**: Generated for any report-like output (headings,
      structured content, lists, tables).
    - **Excel (.xlsx)**: Generated additionally when the output contains
      significant markdown tables (3+ data rows).
    """
    if not _looks_like_report(text, agent_name, subtask_label):
        return []

    results: list[str] = []
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")
    safe_name = re.sub(r"[^a-zA-Z0-9_-]", "", agent_name.replace(" ", "-")).lower()
    # Build filename with subtask label if provided
    safe_subtask = ""
    if subtask_label:
        safe_subtask = re.sub(r"[^a-zA-Z0-9_-]", "", subtask_label.replace(" ", "-")).lower()
    title = f"{subtask_label or agent_name} Report"

    # ── Generate Word document ──────────────────────────────────
    try:
        docx_bytes = _generate_word_doc(text, title)
        blob_name = f"{safe_name}/{timestamp}-{safe_subtask}-report.docx" if safe_subtask else f"{safe_name}/{timestamp}-report.docx"
        result_json = await _upload_binary_blob(
            container,
            blob_name,
            docx_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        results.append(result_json)
        logger.info(
            "doc_generator_word_uploaded",
            agent=agent_name,
            container=container,
            blob=blob_name,
            size=len(docx_bytes),
        )
    except Exception:
        logger.exception("doc_generator_word_failed", agent=agent_name)

    # ── Generate Excel if there are significant tables ──────────
    if _has_significant_tables(text):
        try:
            xlsx_bytes = _generate_excel(text, title)
            blob_name = f"{safe_name}/{timestamp}-{safe_subtask}-report.xlsx" if safe_subtask else f"{safe_name}/{timestamp}-report.xlsx"
            result_json = await _upload_binary_blob(
                container,
                blob_name,
                xlsx_bytes,
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
            results.append(result_json)
            logger.info(
                "doc_generator_excel_uploaded",
                agent=agent_name,
                container=container,
                blob=blob_name,
                size=len(xlsx_bytes),
            )
        except Exception:
            logger.exception("doc_generator_excel_failed", agent=agent_name)

    return results


async def generate_documents_with_summary(
    text: str,
    agent_name: str,
    container: str = "migration-reports",
    subtask_label: str | None = None,
) -> DocumentResult | None:
    """Generate documents and return a summary for chat display.

    Returns a DocumentResult containing:
    - **summary**: Condensed version for chat window display
    - **full_text**: Original detailed text (preserved in document)
    - **uploads**: Blob upload JSON strings for download cards

    Returns None if the text doesn't look like a report worth exporting.
    Skips intermediate subtasks (list, retrieve, fetch, etc.) to avoid
    generating unnecessary documents for data-gathering steps.
    """
    if not _looks_like_report(text, agent_name, subtask_label):
        return None

    # Generate uploads using existing function logic
    uploads = await generate_and_upload_documents(
        text, agent_name, container, subtask_label
    )

    if not uploads:
        return None

    # Generate condensed summary for chat
    summary = _generate_summary(text, agent_name, subtask_label)

    return DocumentResult(
        summary=summary,
        full_text=text,
        uploads=uploads,
    )
